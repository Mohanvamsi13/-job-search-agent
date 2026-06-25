"""Parses a master resume file (.docx or .pdf) into structured ResumeData.

Approach: extract raw text with a deterministic library (no LLM involved in
this step), then make ONE Claude call to organize that exact text into
structured fields. The LLM is told explicitly to copy text, not rephrase it,
so the structured resume stays a faithful representation of the original
file. This is the ground truth fact-checking is performed against later.
"""
import json
from pathlib import Path

from docx import Document as DocxDocument
import pdfplumber

from app.llm_client import complete_json
from app.models import ResumeData

EXTRACTION_SYSTEM_PROMPT = """You convert resume text into structured JSON.

Rules:
- Copy wording from the source resume as closely as possible. Do not add,
  remove, embellish, or infer facts that are not present in the text.
- If a field is not present in the source, leave it empty rather than guessing.
- Preserve numbers, scale, and technology names exactly as written.
- Output ONLY valid JSON matching this schema, with no markdown fences and
  no commentary:

{
  "full_name": "",
  "email": "",
  "phone": "",
  "location": "",
  "nationality": "",
  "linkedin_url": "",
  "summary": "",
  "skills": [""],
  "languages": [""],
  "certifications": [""],
  "experience": [
    {"company": "", "title": "", "start_date": "", "end_date": "",
     "city": "", "country": "", "company_url": "", "location": "", "bullets": [""]}
  ],
  "education": [
    {"institution": "", "degree": "", "field": "", "graduation_date": "",
     "city": "", "country": "", "institution_url": "", "level_eqf": ""}
  ]
}

Notes:
- "nationality": only fill if explicitly stated in the resume (e.g. "Nationality: German"). Leave empty otherwise - never guess from name, address, or any other signal.
- "linkedin_url": only fill if a LinkedIn URL or handle is explicitly present. Leave empty otherwise.
- "languages": only include if the resume explicitly lists language skills
  (e.g. "German (C1), English (Native)", "Languages: German - Fluent").
  Copy proficiency descriptors exactly as written. Leave the list empty if
  no language section exists - do not infer language skills from
  nationality, location, or any other field.
- For each experience entry: "company" should be ONLY the company name
  (e.g. "Reolink"), never combined with city/country. If the source lists
  "City: X | Country: Y" separately, put those in "city" and "country" -
  do not concatenate them into the company or location field. Only fill
  "company_url" if a URL for the employer is explicitly shown. The
  "location" field is a free-text fallback - only use it if the source
  gives an address that doesn't cleanly split into city/country.
- For each education entry: same separation applies for institution vs
  city/country. "level_eqf" should only be filled if the source explicitly
  states an EQF (European Qualifications Framework) level, e.g. "Level in
  EQF: 7" - never infer this from the degree type.
"""


def _extract_text_from_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text)
    return "\n".join(lines)


def _extract_text_from_pdf(path: Path) -> str:
    lines = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            lines.append(text)
    return "\n".join(lines)


def extract_raw_text(file_path: str) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_text_from_docx(path)
    if suffix == ".pdf":
        return _extract_text_from_pdf(path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8")
    raise ValueError(f"Unsupported resume file type: {suffix}")


def structure_resume_text(raw_text: str) -> ResumeData:
    """Single LLM call: raw resume text -> structured ResumeData."""
    text = complete_json(EXTRACTION_SYSTEM_PROMPT, raw_text, max_tokens=4000)
    data = json.loads(text)
    return ResumeData(**data)


def parse_resume(file_path: str) -> ResumeData:
    raw_text = extract_raw_text(file_path)
    if not raw_text.strip():
        raise ValueError("No extractable text found in resume file.")
    return structure_resume_text(raw_text)
