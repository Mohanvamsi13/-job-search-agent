"""Formats a generated cover letter into a German-style formal business
letter (Geschäftsbrief) layout: sender block top-left, date top-right,
recipient block, a bold subject line, then the letter body.

This is purely deterministic formatting - the actual letter content comes
from cover_letter.generate_cover_letter(), already fact-checked. This file
just wraps it in the structural envelope German employers expect, since a
cover letter dropped straight into a blank page (no sender info, no date,
no subject line) reads as informal / non-compliant with German conventions.
"""
from datetime import date

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models import JobDescription, ResumeData
from app.de_format_utils import extract_city


def write_cover_letter_docx(
    resume: ResumeData,
    jd: JobDescription,
    letter_text: str,
    output_path: str,
    signature_city: str = "",
) -> str:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Sender block (top-left) ----
    sender_lines = [resume.full_name, resume.location, resume.phone, resume.email]
    for line in sender_lines:
        if line:
            doc.add_paragraph(line)

    doc.add_paragraph()

    # ---- Date (right-aligned) ----
    city = signature_city or (resume.location.split(",")[0].strip() if resume.location else "")
    today = date.today().strftime("%d.%m.%Y")
    date_p = doc.add_paragraph(f"{city}, {today}" if city else today)
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # ---- Recipient block ----
    if jd.company:
        doc.add_paragraph(jd.company)
    doc.add_paragraph("Attn: Hiring Team")
    if jd.location:
        doc.add_paragraph(jd.location)

    doc.add_paragraph()

    # ---- Subject line ----
    subject_p = doc.add_paragraph()
    subject_text = f"Application for the position of {jd.title}" if jd.title else "Application"
    subject_p.add_run(subject_text).bold = True

    doc.add_paragraph()

    # ---- Body ----
    for para in letter_text.split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)

    doc.save(output_path)
    return output_path
