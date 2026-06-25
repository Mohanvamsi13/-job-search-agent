"""Renders a ResumeData object into a formatted .docx file using python-docx.
No LLM calls here - this is pure deterministic formatting, run only after
you've reviewed and approved the tailored content.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models import ResumeData


def write_resume_docx(resume: ResumeData, output_path: str) -> str:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Header: name + contact line
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(resume.full_name or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(18)

    contact_bits = [b for b in [resume.email, resume.phone, resume.location] if b]
    if contact_bits:
        contact_p = doc.add_paragraph(" | ".join(contact_bits))
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.runs[0].font.size = Pt(9.5)

    if resume.summary:
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(resume.summary)

    if resume.skills:
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(" | ".join(resume.skills))

    if resume.experience:
        doc.add_heading("Experience", level=2)
        for exp in resume.experience:
            title_p = doc.add_paragraph()
            title_run = title_p.add_run(f"{exp.title} - {exp.company}")
            title_run.bold = True
            dates = " - ".join(d for d in [exp.start_date, exp.end_date] if d)
            meta_bits = [b for b in [dates, exp.location] if b]
            if meta_bits:
                title_p.add_run("   " + " | ".join(meta_bits)).italic = True
            for bullet in exp.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    if resume.certifications:
        doc.add_heading("Certifications", level=2)
        doc.add_paragraph(" | ".join(resume.certifications))

    if resume.education:
        doc.add_heading("Education", level=2)
        for edu in resume.education:
            line = " - ".join(b for b in [edu.degree, edu.field, edu.institution] if b)
            doc.add_paragraph(f"{line} ({edu.graduation_date})" if edu.graduation_date else line)

    doc.save(output_path)
    return output_path
