"""Renders a ResumeData object as a German-style tabular CV
(tabellarischer Lebenslauf), following the structural conventions German
employers expect: two-column date/content tables, reverse-chronological
order, a personal details block, a languages section with proficiency
levels, and a closing place/date signature line.

This mirrors the structure used in the EU's Europass CV template without
reproducing Europass's own branded design (logo, exact colors, official
layout grid) - it's an original layout built to the same conventions.

No LLM calls here - pure deterministic formatting of data you already
reviewed. Nothing in this file invents facts: nationality and languages
are only shown if your original resume stated them (see resume_parser.py).
"""
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from app.models import ResumeData
from app.de_format_utils import extract_city

LABELS = {
    "personal": "Persönliche Daten",
    "address": "Adresse",
    "phone": "Telefon",
    "email": "E-Mail",
    "nationality": "Nationalität",
    "profile": "Profil",
    "experience": "Berufserfahrung",
    "education": "Ausbildung",
    "skills": "Kenntnisse & Fähigkeiten",
    "languages": "Sprachkenntnisse",
    "certifications": "Zertifizierungen",
    "present": "heute",
}


def _set_cell_borders_none(cell):
    """python-docx tables get default borders; German CVs typically use a
    clean, borderless two-column layout instead."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.makeelement(qn("w:tcBorders"), {})
    for edge in ("top", "left", "bottom", "right"):
        el = borders.makeelement(qn(f"w:{edge}"), {qn("w:val"): "nil"})
        borders.append(el)
    tc_pr.append(borders)


def _two_col_table(doc, left_width_cm=3.2):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(left_width_cm)
    table.columns[1].width = Cm(15.0 - left_width_cm)
    return table


def _add_row(table, left_text, right_builder):
    """right_builder is a callable that receives the right-hand cell and
    fills it with whatever paragraphs/runs it needs."""
    row = table.add_row()
    left_cell, right_cell = row.cells
    _set_cell_borders_none(left_cell)
    _set_cell_borders_none(right_cell)
    p = left_cell.paragraphs[0]
    p.text = left_text
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9.5)
    right_builder(right_cell)


def write_resume_docx_de(resume: ResumeData, output_path: str, signature_city: str = "") -> str:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ---- Header: name ----
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(resume.full_name or "Name")
    name_run.bold = True
    name_run.font.size = Pt(20)

    # ---- Persönliche Daten ----
    doc.add_paragraph()
    heading = doc.add_paragraph()
    heading.add_run(LABELS["personal"]).bold = True

    personal_table = _two_col_table(doc)
    if resume.location:
        _add_row(personal_table, LABELS["address"], lambda c: c.paragraphs[0].add_run(resume.location))
    if resume.phone:
        _add_row(personal_table, LABELS["phone"], lambda c: c.paragraphs[0].add_run(resume.phone))
    if resume.email:
        _add_row(personal_table, LABELS["email"], lambda c: c.paragraphs[0].add_run(resume.email))
    if resume.nationality:
        _add_row(personal_table, LABELS["nationality"], lambda c: c.paragraphs[0].add_run(resume.nationality))

    # ---- Profil ----
    if resume.summary:
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.add_run(LABELS["profile"]).bold = True
        doc.add_paragraph(resume.summary)

    # ---- Berufserfahrung ----
    if resume.experience:
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.add_run(LABELS["experience"]).bold = True

        exp_table = _two_col_table(doc)
        for exp in resume.experience:
            end = exp.end_date or LABELS["present"]
            date_range = f"{exp.start_date} – {end}".strip(" –")

            def build(cell, exp=exp):
                p = cell.paragraphs[0]
                title_run = p.add_run(f"{exp.title}")
                title_run.bold = True
                p.add_run(f"  |  {exp.company}")
                if exp.location:
                    p.add_run(f", {exp.location}")
                for bullet in exp.bullets:
                    bp = cell.add_paragraph(style="List Bullet")
                    bp.add_run(bullet)

            _add_row(exp_table, date_range, build)

    # ---- Ausbildung ----
    if resume.education:
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.add_run(LABELS["education"]).bold = True

        edu_table = _two_col_table(doc)
        for edu in resume.education:
            def build(cell, edu=edu):
                p = cell.paragraphs[0]
                line = " - ".join(b for b in [edu.degree, edu.field] if b)
                p.add_run(line).bold = True
                if edu.institution:
                    cell.add_paragraph(edu.institution)

            _add_row(edu_table, edu.graduation_date, build)

    # ---- Kenntnisse & Fähigkeiten ----
    if resume.skills:
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.add_run(LABELS["skills"]).bold = True
        doc.add_paragraph(" | ".join(resume.skills))

    # ---- Sprachkenntnisse ----
    if resume.languages:
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.add_run(LABELS["languages"]).bold = True
        doc.add_paragraph(" | ".join(resume.languages))

    # ---- Zertifizierungen ----
    if resume.certifications:
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.add_run(LABELS["certifications"]).bold = True
        doc.add_paragraph(" | ".join(resume.certifications))

    # ---- Signature line ----
    doc.add_paragraph()
    doc.add_paragraph()
    city = signature_city or extract_city(resume.location)
    today = date.today().strftime("%d.%m.%Y")
    sig_line = f"{city}, {today}" if city else f"___________, {today}"
    doc.add_paragraph(sig_line)

    doc.save(output_path)
    return output_path
